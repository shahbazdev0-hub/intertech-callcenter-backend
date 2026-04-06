#  # backend/app/services/rag_service.py - RAG SERVICE FOR DOCUMENT TRAINING
# # ✅ ENHANCED: Added Dynamic Summary Context Generation

# import os
# import logging
# import hashlib
# from datetime import datetime
# from typing import List, Dict, Any, Optional, Tuple
# from pathlib import Path
# from motor.motor_asyncio import AsyncIOMotorDatabase
# from bson import ObjectId
# import aiofiles
# import numpy as np

# # Document processing
# import PyPDF2
# from docx import Document as DocxDocument
# import io

# # OpenAI for embeddings
# import openai
# from openai import AsyncOpenAI

# logger = logging.getLogger(__name__)


# class RAGService:
#     """
#     Retrieval-Augmented Generation Service
#     Handles document upload, processing, embedding generation, and similarity search
    
#     ✅ ENHANCED: Now includes Dynamic Summary Context Generation
#     """
    
#     def __init__(self):
#         self.openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
#         self.upload_base_path = Path("uploads/agent_documents")
#         self.upload_base_path.mkdir(parents=True, exist_ok=True)
        
#         # RAG Configuration - ✅ OPTIMIZED FOR VOICE
#         self.chunk_size = int(os.getenv("RAG_CHUNK_SIZE", 500))
#         self.chunk_overlap = int(os.getenv("RAG_CHUNK_OVERLAP", 50))
#         self.embedding_model = os.getenv("RAG_EMBEDDING_MODEL", "text-embedding-ada-002")
        
#         # ✅ FIX: Lower threshold for voice conversations
#         # Voice queries are informal and may not match document text exactly
#         self.similarity_threshold = float(os.getenv("RAG_SIMILARITY_THRESHOLD", 0.65))
#         self.max_results = int(os.getenv("RAG_MAX_RESULTS", 3))
        
#         logger.info(f"📚 RAG Service initialized:")
#         logger.info(f"   Chunk size: {self.chunk_size}")
#         logger.info(f"   Similarity threshold: {self.similarity_threshold}")
#         logger.info(f"   Max results: {self.max_results}")
    
    
#     # ============================================
#     # ✅ NEW: GENERATE AGENT CONTEXT (DYNAMIC SUMMARY)
#     # ============================================
    
#     async def generate_agent_context(
#         self,
#         agent_id: str,
#         user_id: str,
#         db: AsyncIOMotorDatabase,
#         script_text: Optional[str] = None,
#         force_regenerate: bool = False
#     ) -> Dict[str, Any]:
#         """
#         ✅ NEW: Generate comprehensive agent context from documents and script
        
#         This creates a pre-built summary that will be injected into every call
#         for instant contextual responses (no RAG search needed during calls).
        
#         Args:
#             agent_id: The agent's ID
#             user_id: The user's ID
#             db: Database connection
#             script_text: Custom AI script (optional, will fetch from agent if not provided)
#             force_regenerate: Force regeneration even if context exists
            
#         Returns:
#             Dict with success status and generated context
#         """
#         try:
#             logger.info(f"\n{'='*60}")
#             logger.info(f"🧠 GENERATING AGENT CONTEXT")
#             logger.info(f"{'='*60}")
#             logger.info(f"   Agent ID: {agent_id}")
#             logger.info(f"   Force Regenerate: {force_regenerate}")
#             logger.info(f"{'='*60}\n")
            
#             # Import OpenAI service
#             from .openai import openai_service
            
#             # Get agent data
#             agent = await db.voice_agents.find_one({
#                 "_id": ObjectId(agent_id),
#                 "user_id": user_id
#             })
            
#             if not agent:
#                 logger.error(f"❌ Agent not found: {agent_id}")
#                 return {
#                     "success": False,
#                     "error": "Agent not found"
#                 }
            
#             # Check if context already exists and we're not forcing regeneration
#             existing_context = agent.get("agent_context")
#             if existing_context and not force_regenerate:
#                 logger.info(f"✅ Using existing agent context (generated at: {existing_context.get('generated_at')})")
#                 return {
#                     "success": True,
#                     "context": existing_context,
#                     "cached": True
#                 }
            
#             # Get script text from agent if not provided
#             if not script_text:
#                 script_text = agent.get("ai_script") or agent.get("system_prompt") or ""
            
#             # Get all document texts for this agent
#             document_text = await self._get_all_document_text(agent_id, db)
            
#             logger.info(f"📄 Collected data:")
#             logger.info(f"   Script length: {len(script_text)} chars")
#             logger.info(f"   Document text length: {len(document_text)} chars")
            
#             # Check if we have anything to process
#             if not script_text and not document_text:
#                 logger.warning(f"⚠️ No script or documents found for agent {agent_id}")
#                 return {
#                     "success": False,
#                     "error": "No script or documents to generate context from"
#                 }
            
#             # Generate context using OpenAI
#             summary_result = await openai_service.generate_agent_summary(
#                 document_text=document_text if document_text else None,
#                 script_text=script_text if script_text else None,
#                 existing_context=existing_context if not force_regenerate else None
#             )
            
#             if not summary_result.get("success"):
#                 logger.error(f"❌ Failed to generate summary: {summary_result.get('error')}")
#                 return {
#                     "success": False,
#                     "error": summary_result.get("error", "Summary generation failed")
#                 }
            
#             # Get the generated context
#             context = summary_result.get("context", {})
            
#             # Add metadata
#             context["generated_at"] = datetime.utcnow().isoformat()
#             context["script_included"] = bool(script_text)
#             context["documents_included"] = bool(document_text)
            
#             # Get source document names
#             source_docs = await db.agent_documents.find(
#                 {"agent_id": agent_id, "processing_status": "completed"},
#                 {"filename": 1}
#             ).to_list(length=None)
#             context["source_documents"] = [doc.get("filename", "") for doc in source_docs]
            
#             # Store context in agent record
#             await db.voice_agents.update_one(
#                 {"_id": ObjectId(agent_id)},
#                 {
#                     "$set": {
#                         "agent_context": context,
#                         "has_context": True,
#                         "context_generated_at": datetime.utcnow(),
#                         "updated_at": datetime.utcnow()
#                     }
#                 }
#             )
            
#             logger.info(f"✅ Agent context generated and stored successfully")
#             logger.info(f"   Identity: {context.get('identity', {})}")
#             logger.info(f"   Services: {len(context.get('company_info', {}).get('services', []))} items")
#             logger.info(f"   FAQs: {len(context.get('faqs', []))} items")
#             logger.info(f"   Source documents: {context.get('source_documents', [])}")
            
#             return {
#                 "success": True,
#                 "context": context,
#                 "cached": False
#             }
            
#         except Exception as e:
#             logger.error(f"❌ Error generating agent context: {e}")
#             import traceback
#             traceback.print_exc()
#             return {
#                 "success": False,
#                 "error": str(e)
#             }
    
    
#     async def _get_all_document_text(
#         self,
#         agent_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> str:
#         """
#         Get all extracted text from agent's documents
        
#         Args:
#             agent_id: The agent's ID
#             db: Database connection
            
#         Returns:
#             Combined text from all documents
#         """
#         try:
#             # Get all completed documents for this agent
#             documents = await db.agent_documents.find({
#                 "agent_id": agent_id,
#                 "processing_status": "completed"
#             }).to_list(length=None)
            
#             if not documents:
#                 logger.info(f"ℹ️ No documents found for agent {agent_id}")
#                 return ""
            
#             # Combine all document texts
#             all_texts = []
#             for doc in documents:
#                 extracted_text = doc.get("extracted_text", "")
#                 if extracted_text:
#                     filename = doc.get("filename", "Unknown")
#                     all_texts.append(f"--- Document: {filename} ---\n{extracted_text}")
            
#             combined_text = "\n\n".join(all_texts)
#             logger.info(f"📄 Combined {len(documents)} documents ({len(combined_text)} chars)")
            
#             return combined_text
            
#         except Exception as e:
#             logger.error(f"❌ Error getting document text: {e}")
#             return ""
    
    
#     async def update_agent_context_on_script_change(
#         self,
#         agent_id: str,
#         user_id: str,
#         new_script: str,
#         db: AsyncIOMotorDatabase
#     ) -> Dict[str, Any]:
#         """
#         ✅ NEW: Update agent context when script changes
        
#         Args:
#             agent_id: The agent's ID
#             user_id: The user's ID
#             new_script: The new AI script
#             db: Database connection
            
#         Returns:
#             Dict with success status
#         """
#         try:
#             logger.info(f"🔄 Updating agent context for script change (Agent: {agent_id})")
            
#             # Regenerate context with new script
#             result = await self.generate_agent_context(
#                 agent_id=agent_id,
#                 user_id=user_id,
#                 db=db,
#                 script_text=new_script,
#                 force_regenerate=True
#             )
            
#             return result
            
#         except Exception as e:
#             logger.error(f"❌ Error updating context on script change: {e}")
#             return {
#                 "success": False,
#                 "error": str(e)
#             }
    
    
#     async def update_agent_context_on_document_change(
#         self,
#         agent_id: str,
#         user_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> Dict[str, Any]:
#         """
#         ✅ NEW: Update agent context when documents change (upload/delete)
        
#         Args:
#             agent_id: The agent's ID
#             user_id: The user's ID
#             db: Database connection
            
#         Returns:
#             Dict with success status
#         """
#         try:
#             logger.info(f"🔄 Updating agent context for document change (Agent: {agent_id})")
            
#             # Regenerate context with current documents
#             result = await self.generate_agent_context(
#                 agent_id=agent_id,
#                 user_id=user_id,
#                 db=db,
#                 force_regenerate=True
#             )
            
#             return result
            
#         except Exception as e:
#             logger.error(f"❌ Error updating context on document change: {e}")
#             return {
#                 "success": False,
#                 "error": str(e)
#             }
    
    
#     async def get_agent_context(
#         self,
#         agent_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> Optional[Dict[str, Any]]:
#         """
#         ✅ NEW: Get pre-generated agent context (fast lookup)
        
#         Args:
#             agent_id: The agent's ID
#             db: Database connection
            
#         Returns:
#             Agent context dict or None
#         """
#         try:
#             agent = await db.voice_agents.find_one(
#                 {"_id": ObjectId(agent_id)},
#                 {"agent_context": 1, "has_context": 1}
#             )
            
#             if agent and agent.get("has_context") and agent.get("agent_context"):
#                 return agent.get("agent_context")
            
#             return None
            
#         except Exception as e:
#             logger.error(f"❌ Error getting agent context: {e}")
#             return None
    
    
#     # ============================================
#     # DOCUMENT UPLOAD & PROCESSING
#     # ============================================
    
#     async def upload_and_process_document(
#         self,
#         file_content: bytes,
#         filename: str,
#         content_type: str,
#         agent_id: str,
#         user_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> Dict[str, Any]:
#         """
#         Upload document and process it for RAG
        
#         ✅ ENHANCED: Now also triggers context regeneration after upload
        
#         Steps:
#         1. Validate file
#         2. Save to disk
#         3. Extract text
#         4. Create chunks
#         5. Generate embeddings
#         6. Store in database
#         7. ✅ NEW: Regenerate agent context
#         """
#         try:
#             logger.info(f"📄 Starting document upload for agent {agent_id}")
            
#             # Step 1: Validate file
#             validation = self._validate_file(filename, content_type, len(file_content))
#             if not validation['valid']:
#                 return {
#                     "success": False,
#                     "error": validation['error']
#                 }
            
#             # Step 2: Save file to disk
#             file_path = await self._save_file(
#                 file_content,
#                 filename,
#                 agent_id,
#                 user_id
#             )
            
#             # Step 3: Create document record
#             doc_id = await self._create_document_record(
#                 agent_id=agent_id,
#                 user_id=user_id,
#                 filename=filename,
#                 file_path=str(file_path),
#                 file_type=content_type,
#                 file_size=len(file_content),
#                 db=db
#             )
            
#             # Step 4: Extract text from document
#             logger.info(f"📖 Extracting text from {filename}")
#             extracted_text = await self._extract_text(
#                 file_content,
#                 content_type
#             )
            
#             if not extracted_text:
#                 await self._update_document_status(
#                     doc_id,
#                     "failed",
#                     "Failed to extract text from document",
#                     db
#                 )
#                 return {
#                     "success": False,
#                     "error": "Failed to extract text from document"
#                 }
            
#             # Step 5: Create text chunks
#             logger.info(f"✂️ Creating text chunks")
#             chunks = self._create_chunks(extracted_text)
            
#             # Step 6: Generate embeddings for chunks
#             logger.info(f"🧠 Generating embeddings for {len(chunks)} chunks")
#             chunks_with_embeddings = await self._generate_embeddings(chunks)
            
#             # Step 7: Store chunks with embeddings in database
#             await self._store_chunks(
#                 doc_id=doc_id,
#                 agent_id=agent_id,
#                 chunks=chunks_with_embeddings,
#                 extracted_text=extracted_text,
#                 db=db
#             )
            
#             # Step 8: Update agent's has_training_docs flag
#             await self._update_agent_training_status(agent_id, doc_id, db)
            
#             logger.info(f"✅ Document processed successfully: {filename}")
            
#             # ✅ NEW: Step 9 - Regenerate agent context with new document
#             logger.info(f"🔄 Regenerating agent context after document upload...")
#             context_result = await self.update_agent_context_on_document_change(
#                 agent_id=agent_id,
#                 user_id=user_id,
#                 db=db
#             )
            
#             if context_result.get("success"):
#                 logger.info(f"✅ Agent context updated with new document")
#             else:
#                 logger.warning(f"⚠️ Failed to update context: {context_result.get('error')}")
            
#             return {
#                 "success": True,
#                 "document_id": doc_id,
#                 "filename": filename,
#                 "file_size": len(file_content),
#                 "total_chunks": len(chunks_with_embeddings),
#                 "processing_status": "completed",
#                 "context_updated": context_result.get("success", False)
#             }
            
#         except Exception as e:
#             logger.error(f"❌ Error processing document: {e}", exc_info=True)
#             return {
#                 "success": False,
#                 "error": str(e)
#             }
    
    
#     # ============================================
#     # TEXT EXTRACTION
#     # ============================================
    
#     async def _extract_text(
#         self,
#         file_content: bytes,
#         content_type: str
#     ) -> Optional[str]:
#         """Extract text from different file types"""
#         try:
#             if content_type == "application/pdf":
#                 return await self._extract_text_from_pdf(file_content)
            
#             elif content_type in [
#                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 "application/msword"
#             ]:
#                 return await self._extract_text_from_docx(file_content)
            
#             elif content_type in ["text/plain", "text/markdown"]:
#                 return file_content.decode('utf-8', errors='ignore')
            
#             else:
#                 logger.warning(f"Unsupported content type: {content_type}")
#                 return None
                
#         except Exception as e:
#             logger.error(f"Error extracting text: {e}")
#             return None
    
    
#     async def _extract_text_from_pdf(self, file_content: bytes) -> Optional[str]:
#         """Extract text from PDF file"""
#         try:
#             pdf_file = io.BytesIO(file_content)
#             pdf_reader = PyPDF2.PdfReader(pdf_file)
            
#             text_parts = []
#             for page in pdf_reader.pages:
#                 text = page.extract_text()
#                 if text:
#                     text_parts.append(text)
            
#             return "\n\n".join(text_parts)
            
#         except Exception as e:
#             logger.error(f"Error extracting PDF text: {e}")
#             return None
    
    
#     async def _extract_text_from_docx(self, file_content: bytes) -> Optional[str]:
#         """Extract text from DOCX file"""
#         try:
#             docx_file = io.BytesIO(file_content)
#             doc = DocxDocument(docx_file)
            
#             text_parts = []
#             for paragraph in doc.paragraphs:
#                 if paragraph.text.strip():
#                     text_parts.append(paragraph.text)
            
#             # Also extract from tables
#             for table in doc.tables:
#                 for row in table.rows:
#                     row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
#                     if row_text:
#                         text_parts.append(" | ".join(row_text))
            
#             return "\n\n".join(text_parts)
            
#         except Exception as e:
#             logger.error(f"Error extracting DOCX text: {e}")
#             return None
    
    
#     # ============================================
#     # TEXT CHUNKING
#     # ============================================
    
#     def _create_chunks(self, text: str) -> List[str]:
#         """Split text into overlapping chunks"""
#         if not text:
#             return []
        
#         # Clean text
#         text = text.strip()
        
#         # Split into sentences first
#         sentences = self._split_into_sentences(text)
        
#         chunks = []
#         current_chunk = []
#         current_length = 0
        
#         for sentence in sentences:
#             sentence_length = len(sentence)
            
#             if current_length + sentence_length > self.chunk_size and current_chunk:
#                 # Save current chunk
#                 chunks.append(" ".join(current_chunk))
                
#                 # Start new chunk with overlap
#                 overlap_start = max(0, len(current_chunk) - 2)
#                 current_chunk = current_chunk[overlap_start:]
#                 current_length = sum(len(s) for s in current_chunk)
            
#             current_chunk.append(sentence)
#             current_length += sentence_length
        
#         # Don't forget the last chunk
#         if current_chunk:
#             chunks.append(" ".join(current_chunk))
        
#         logger.info(f"Created {len(chunks)} chunks from {len(text)} chars")
#         return chunks
    
    
#     def _split_into_sentences(self, text: str) -> List[str]:
#         """Split text into sentences"""
#         import re
        
#         # Simple sentence splitting
#         sentences = re.split(r'(?<=[.!?])\s+', text)
#         sentences = [s.strip() for s in sentences if s.strip()]
        
#         return sentences
    
    
#     # ============================================
#     # EMBEDDING GENERATION
#     # ============================================
    
#     async def _generate_embeddings(
#         self,
#         chunks: List[str]
#     ) -> List[Dict[str, Any]]:
#         """Generate OpenAI embeddings for text chunks"""
#         try:
#             chunks_with_embeddings = []
            
#             # Process chunks in batches to avoid rate limits
#             batch_size = 10
#             for i in range(0, len(chunks), batch_size):
#                 batch = chunks[i:i + batch_size]
                
#                 # Call OpenAI API
#                 response = await self.openai_client.embeddings.create(
#                     model=self.embedding_model,
#                     input=batch
#                 )
                
#                 # Extract embeddings
#                 for idx, chunk_text in enumerate(batch):
#                     embedding = response.data[idx].embedding
                    
#                     chunks_with_embeddings.append({
#                         "chunk_id": f"chunk_{i + idx}",
#                         "text": chunk_text,
#                         "embedding": embedding,
#                         "chunk_index": i + idx
#                     })
            
#             logger.info(f"Generated {len(chunks_with_embeddings)} embeddings")
#             return chunks_with_embeddings
            
#         except Exception as e:
#             logger.error(f"Error generating embeddings: {e}")
#             raise
    
    
#     # ============================================
#     # SIMILARITY SEARCH (KEPT FOR BACKWARD COMPATIBILITY)
#     # ============================================
    
#     async def query_documents(
#         self,
#         agent_id: str,
#         query: str,
#         db: AsyncIOMotorDatabase,
#         top_k: int = 3
#     ) -> Dict[str, Any]:
#         """
#         Search agent's documents for relevant information
        
#         ✅ NOTE: This is kept for backward compatibility.
#         The new approach uses pre-generated context instead of real-time search.
        
#         Returns:
#         - relevant chunks
#         - confidence score
#         - combined answer
#         """
#         try:
#             logger.info(f"\n{'='*60}")
#             logger.info(f"🔍 RAG DOCUMENT SEARCH")
#             logger.info(f"{'='*60}")
#             logger.info(f"   Agent ID: {agent_id}")
#             logger.info(f"   Query: '{query}'")
#             logger.info(f"{'='*60}\n")
            
#             # Step 1: Get agent's documents
#             documents = await db.agent_documents.find({
#                 "agent_id": agent_id,
#                 "processing_status": "completed"
#             }).to_list(length=None)
            
#             logger.info(f"📄 Found {len(documents)} documents for agent {agent_id}")
            
#             if not documents:
#                 # Check if documents exist with different status
#                 all_docs = await db.agent_documents.find({
#                     "agent_id": agent_id
#                 }).to_list(length=None)
                
#                 if all_docs:
#                     logger.warning(f"⚠️ Agent has {len(all_docs)} documents but none are 'completed'")
#                     for doc in all_docs:
#                         logger.warning(f"   - {doc.get('filename')}: status={doc.get('processing_status')}")
#                 else:
#                     logger.info(f"ℹ️ No documents found for agent {agent_id}")
                
#                 return {
#                     "found": False,
#                     "confidence": 0.0,
#                     "answer": None,
#                     "sources": []
#                 }
            
#             # Log document details
#             for doc in documents:
#                 chunk_count = len(doc.get("chunks", []))
#                 logger.info(f"   📄 {doc.get('filename')}: {chunk_count} chunks")
            
#             # Step 2: Generate query embedding
#             logger.info(f"🧠 Generating query embedding...")
#             query_embedding = await self._generate_query_embedding(query)
#             logger.info(f"✅ Query embedding generated ({len(query_embedding)} dimensions)")
            
#             # Step 3: Search all chunks for similarity
#             all_matches = []
#             search_threshold = self.similarity_threshold - 0.1  # Lower threshold for voice
            
#             for doc in documents:
#                 chunks = doc.get("chunks", [])
#                 for chunk in chunks:
#                     chunk_embedding = chunk.get("embedding", [])
#                     if chunk_embedding:
#                         similarity = self._cosine_similarity(query_embedding, chunk_embedding)
                        
#                         if similarity >= search_threshold:
#                             all_matches.append({
#                                 "text": chunk.get("text", ""),
#                                 "similarity": similarity,
#                                 "document_id": str(doc["_id"]),
#                                 "filename": doc.get("filename", "Unknown"),
#                                 "chunk_id": chunk.get("chunk_id", "")
#                             })
            
#             logger.info(f"📊 Found {len(all_matches)} matches above threshold {search_threshold}")
            
#             # Step 4: Sort by similarity and get top K
#             all_matches.sort(key=lambda x: x["similarity"], reverse=True)
#             top_matches = all_matches[:top_k]
            
#             if not top_matches:
#                 logger.info(f"❌ No relevant matches found above threshold {search_threshold}")
#                 return {
#                     "found": False,
#                     "confidence": 0.0,
#                     "answer": None,
#                     "sources": []
#                 }
            
#             # Step 5: Combine top matches into answer
#             avg_confidence = sum(m["similarity"] for m in top_matches) / len(top_matches)
#             combined_text = "\n\n".join([m["text"] for m in top_matches])
            
#             logger.info(f"\n✅ RAG RESULTS:")
#             logger.info(f"   Top matches: {len(top_matches)}")
#             logger.info(f"   Avg confidence: {avg_confidence:.2f}")
#             for i, match in enumerate(top_matches):
#                 logger.info(f"   Match {i+1}: {match['similarity']:.3f} from {match['filename']}")
#                 logger.info(f"            '{match['text'][:80]}...'")
            
#             # Step 6: Update usage stats
#             for match in top_matches:
#                 await db.agent_documents.update_one(
#                     {"_id": ObjectId(match["document_id"])},
#                     {
#                         "$inc": {"usage_count": 1},
#                         "$set": {"last_used": datetime.utcnow()}
#                     }
#                 )
            
#             return {
#                 "found": True,
#                 "confidence": avg_confidence,
#                 "answer": combined_text,
#                 "sources": top_matches,
#                 "total_matches": len(all_matches)
#             }
            
#         except Exception as e:
#             logger.error(f"❌ Error querying documents: {e}", exc_info=True)
#             return {
#                 "found": False,
#                 "confidence": 0.0,
#                 "answer": None,
#                 "error": str(e)
#             }
    
    
#     async def _generate_query_embedding(self, query: str) -> List[float]:
#         """Generate embedding for search query"""
#         try:
#             response = await self.openai_client.embeddings.create(
#                 model=self.embedding_model,
#                 input=[query]
#             )
#             return response.data[0].embedding
#         except Exception as e:
#             logger.error(f"Error generating query embedding: {e}")
#             raise
    
    
#     def _cosine_similarity(
#         self,
#         vec1: List[float],
#         vec2: List[float]
#     ) -> float:
#         """Calculate cosine similarity between two vectors"""
#         try:
#             vec1_np = np.array(vec1)
#             vec2_np = np.array(vec2)
            
#             dot_product = np.dot(vec1_np, vec2_np)
#             norm1 = np.linalg.norm(vec1_np)
#             norm2 = np.linalg.norm(vec2_np)
            
#             if norm1 == 0 or norm2 == 0:
#                 return 0.0
            
#             return float(dot_product / (norm1 * norm2))
#         except Exception as e:
#             logger.error(f"Error calculating cosine similarity: {e}")
#             return 0.0
    
    
#     # ============================================
#     # FILE MANAGEMENT
#     # ============================================
    
#     def _validate_file(
#         self,
#         filename: str,
#         content_type: str,
#         file_size: int
#     ) -> Dict[str, Any]:
#         """Validate uploaded file"""
        
#         # Check file size (max 10MB)
#         max_size = 10 * 1024 * 1024  # 10MB
#         if file_size > max_size:
#             return {
#                 "valid": False,
#                 "error": f"File size exceeds maximum allowed size of 10MB (got {file_size / 1024 / 1024:.2f}MB)"
#             }
        
#         # Check file type
#         allowed_types = [
#             "application/pdf",
#             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             "application/msword",
#             "text/plain",
#             "text/markdown"
#         ]
        
#         if content_type not in allowed_types:
#             return {
#                 "valid": False,
#                 "error": f"Unsupported file type: {content_type}. Allowed: PDF, DOCX, TXT, MD"
#             }
        
#         return {"valid": True}
    
    
#     async def _save_file(
#         self,
#         file_content: bytes,
#         filename: str,
#         agent_id: str,
#         user_id: str
#     ) -> Path:
#         """Save uploaded file to disk"""
#         try:
#             # Create directory structure: uploads/agent_documents/{user_id}/{agent_id}/
#             agent_dir = self.upload_base_path / user_id / agent_id
#             agent_dir.mkdir(parents=True, exist_ok=True)
            
#             # Generate unique filename with hash
#             file_hash = hashlib.md5(file_content).hexdigest()[:8]
#             safe_filename = f"{file_hash}_{filename}"
            
#             file_path = agent_dir / safe_filename
            
#             # Write file asynchronously
#             async with aiofiles.open(file_path, 'wb') as f:
#                 await f.write(file_content)
            
#             logger.info(f"💾 Saved file to {file_path}")
#             return file_path
            
#         except Exception as e:
#             logger.error(f"Error saving file: {e}")
#             raise
    
    
#     async def delete_document(
#         self,
#         document_id: str,
#         agent_id: str,
#         user_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> Dict[str, Any]:
#         """
#         Delete document and its file
        
#         ✅ ENHANCED: Now also triggers context regeneration after delete
#         """
#         try:
#             # Get document record
#             doc = await db.agent_documents.find_one({
#                 "_id": ObjectId(document_id),
#                 "agent_id": agent_id,
#                 "user_id": user_id
#             })
            
#             if not doc:
#                 return {
#                     "success": False,
#                     "error": "Document not found"
#                 }
            
#             # Delete file from disk
#             file_path = Path(doc["file_path"])
#             if file_path.exists():
#                 file_path.unlink()
#                 logger.info(f"🗑️ Deleted file: {file_path}")
            
#             # Delete document record
#             await db.agent_documents.delete_one({"_id": ObjectId(document_id)})
            
#             # Update agent's training doc list
#             await db.voice_agents.update_one(
#                 {"_id": ObjectId(agent_id)},
#                 {
#                     "$pull": {"training_doc_ids": document_id},
#                     "$set": {"updated_at": datetime.utcnow()}
#                 }
#             )
            
#             # Check if agent still has documents
#             remaining_docs = await db.agent_documents.count_documents({
#                 "agent_id": agent_id,
#                 "processing_status": "completed"
#             })
            
#             if remaining_docs == 0:
#                 await db.voice_agents.update_one(
#                     {"_id": ObjectId(agent_id)},
#                     {"$set": {"has_training_docs": False}}
#                 )
            
#             logger.info(f"✅ Document deleted: {document_id}")
            
#             # ✅ NEW: Regenerate agent context after document deletion
#             logger.info(f"🔄 Regenerating agent context after document deletion...")
#             context_result = await self.update_agent_context_on_document_change(
#                 agent_id=agent_id,
#                 user_id=user_id,
#                 db=db
#             )
            
#             if context_result.get("success"):
#                 logger.info(f"✅ Agent context updated after document deletion")
#             else:
#                 logger.warning(f"⚠️ Failed to update context: {context_result.get('error')}")
            
#             return {
#                 "success": True,
#                 "message": "Document deleted successfully",
#                 "context_updated": context_result.get("success", False)
#             }
            
#         except Exception as e:
#             logger.error(f"❌ Error deleting document: {e}")
#             return {
#                 "success": False,
#                 "error": str(e)
#             }
    
    
#     async def _create_document_record(
#         self,
#         agent_id: str,
#         user_id: str,
#         filename: str,
#         file_path: str,
#         file_type: str,
#         file_size: int,
#         db: AsyncIOMotorDatabase
#     ) -> str:
#         """Create initial document record in database"""
#         doc_data = {
#             "agent_id": agent_id,
#             "user_id": user_id,
#             "filename": filename,
#             "file_path": file_path,
#             "file_type": file_type,
#             "file_size": file_size,
#             "processing_status": "processing",
#             "processing_started_at": datetime.utcnow(),
#             "chunks": [],
#             "total_chunks": 0,
#             "usage_count": 0,
#             "upload_date": datetime.utcnow(),
#             "created_at": datetime.utcnow(),
#             "updated_at": datetime.utcnow()
#         }
        
#         result = await db.agent_documents.insert_one(doc_data)
#         return str(result.inserted_id)
    
    
#     async def _store_chunks(
#         self,
#         doc_id: str,
#         agent_id: str,
#         chunks: List[Dict[str, Any]],
#         extracted_text: str,
#         db: AsyncIOMotorDatabase
#     ):
#         """Store processed chunks in database"""
#         await db.agent_documents.update_one(
#             {"_id": ObjectId(doc_id)},
#             {
#                 "$set": {
#                     "chunks": chunks,
#                     "total_chunks": len(chunks),
#                     "extracted_text": extracted_text[:10000],  # Store first 10K chars
#                     "processing_status": "completed",
#                     "processing_completed_at": datetime.utcnow(),
#                     "updated_at": datetime.utcnow(),
#                     "metadata": {
#                         "word_count": len(extracted_text.split()),
#                         "char_count": len(extracted_text)
#                     }
#                 }
#             }
#         )
    
    
#     async def _update_document_status(
#         self,
#         doc_id: str,
#         status: str,
#         error_message: Optional[str],
#         db: AsyncIOMotorDatabase
#     ):
#         """Update document processing status"""
#         update_data = {
#             "processing_status": status,
#             "updated_at": datetime.utcnow()
#         }
        
#         if error_message:
#             update_data["processing_error"] = error_message
        
#         await db.agent_documents.update_one(
#             {"_id": ObjectId(doc_id)},
#             {"$set": update_data}
#         )
    
    
#     async def _update_agent_training_status(
#         self,
#         agent_id: str,
#         doc_id: str,
#         db: AsyncIOMotorDatabase
#     ):
#         """Update agent's training document flags"""
#         await db.voice_agents.update_one(
#             {"_id": ObjectId(agent_id)},
#             {
#                 "$set": {
#                     "has_training_docs": True,
#                     "updated_at": datetime.utcnow()
#                 },
#                 "$addToSet": {
#                     "training_doc_ids": doc_id
#                 }
#             }
#         )
    
    
#     # ============================================
#     # UTILITY METHODS
#     # ============================================
    
#     async def get_agent_documents(
#         self,
#         agent_id: str,
#         user_id: str,
#         db: AsyncIOMotorDatabase
#     ) -> List[Dict[str, Any]]:
#         """Get all documents for an agent"""
#         try:
#             documents = await db.agent_documents.find({
#                 "agent_id": agent_id,
#                 "user_id": user_id
#             }).to_list(length=None)
            
#             # Convert ObjectId to string
#             for doc in documents:
#                 doc["_id"] = str(doc["_id"])
#                 doc["upload_date"] = doc["upload_date"].isoformat()
                
#                 # Remove embeddings from response (too large)
#                 if "chunks" in doc:
#                     for chunk in doc["chunks"]:
#                         chunk.pop("embedding", None)
            
#             return documents
            
#         except Exception as e:
#             logger.error(f"Error getting agent documents: {e}")
#             return []


# # Create singleton instance
# rag_service = RAGService()


 # backend/app/services/rag_service.py - RAG SERVICE FOR DOCUMENT TRAINING
# ✅ ENHANCED: Added Dynamic Summary Context Generation

import os
import logging
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from motor.motor_asyncio import AsyncIOMotorDatabase
from bson import ObjectId
import aiofiles
import numpy as np

# Document processing
import PyPDF2
from docx import Document as DocxDocument
import io

# OpenAI for embeddings
import openai
from openai import AsyncOpenAI

logger = logging.getLogger(__name__)


class RAGService:
    """
    Retrieval-Augmented Generation Service
    Handles document upload, processing, embedding generation, and similarity search
    
    ✅ ENHANCED: Now includes Dynamic Summary Context Generation
    """
    
    def __init__(self):
        # ✅ FIX: Handle missing API key gracefully
        api_key = os.getenv("OPENAI_API_KEY")
        
        if api_key:
            self.openai_client = AsyncOpenAI(api_key=api_key)
            self.configured = True
            logger.info("✅ RAG Service: OpenAI configured")
        else:
            self.openai_client = None
            self.configured = False
            logger.warning("⚠️ RAG Service: OpenAI API key not configured")
        
        self.upload_base_path = Path("uploads/agent_documents")
        self.upload_base_path.mkdir(parents=True, exist_ok=True)
        
        # RAG Configuration - ✅ OPTIMIZED FOR VOICE
        self.chunk_size = int(os.getenv("RAG_CHUNK_SIZE", 500))
        self.chunk_overlap = int(os.getenv("RAG_CHUNK_OVERLAP", 50))
        self.embedding_model = os.getenv("RAG_EMBEDDING_MODEL", "text-embedding-ada-002")
        
        # ✅ FIX: Lower threshold for voice conversations
        # Voice queries are informal and may not match document text exactly
        self.similarity_threshold = float(os.getenv("RAG_SIMILARITY_THRESHOLD", 0.65))
        self.max_results = int(os.getenv("RAG_MAX_RESULTS", 3))
        
        logger.info(f"📚 RAG Service initialized:")
        logger.info(f"   Chunk size: {self.chunk_size}")
        logger.info(f"   Similarity threshold: {self.similarity_threshold}")
        logger.info(f"   Max results: {self.max_results}")
    
    
    # ============================================
    # ✅ NEW: GENERATE AGENT CONTEXT (DYNAMIC SUMMARY)
    # ============================================
    
    async def generate_agent_context(
        self,
        agent_id: str,
        user_id: str,
        db: AsyncIOMotorDatabase,
        script_text: Optional[str] = None,
        force_regenerate: bool = False
    ) -> Dict[str, Any]:
        """
        ✅ NEW: Generate comprehensive agent context from documents and script
        
        This creates a pre-built summary that will be injected into every call
        for instant contextual responses (no RAG search needed during calls).
        
        Args:
            agent_id: The agent's ID
            user_id: The user's ID
            db: Database connection
            script_text: Custom AI script (optional, will fetch from agent if not provided)
            force_regenerate: Force regeneration even if context exists
            
        Returns:
            Dict with success status and generated context
        """
        try:
            logger.info(f"\n{'='*60}")
            logger.info(f"🧠 GENERATING AGENT CONTEXT")
            logger.info(f"{'='*60}")
            logger.info(f"   Agent ID: {agent_id}")
            logger.info(f"   Force Regenerate: {force_regenerate}")
            logger.info(f"{'='*60}\n")
            
            # Import OpenAI service
            from .openai import openai_service
            
            # Get agent data
            agent = await db.voice_agents.find_one({
                "_id": ObjectId(agent_id),
                "user_id": user_id
            })
            
            if not agent:
                logger.error(f"❌ Agent not found: {agent_id}")
                return {
                    "success": False,
                    "error": "Agent not found"
                }
            
            # Check if context already exists and we're not forcing regeneration
            existing_context = agent.get("agent_context")
            if existing_context and not force_regenerate:
                logger.info(f"✅ Using existing agent context (generated at: {existing_context.get('generated_at')})")
                return {
                    "success": True,
                    "context": existing_context,
                    "cached": True
                }
            
            # Get script text from agent if not provided
            if not script_text:
                script_text = agent.get("ai_script") or agent.get("system_prompt") or ""
            
            # Get all document texts for this agent
            document_text = await self._get_all_document_text(agent_id, db)
            
            logger.info(f"📄 Collected data:")
            logger.info(f"   Script length: {len(script_text)} chars")
            logger.info(f"   Document text length: {len(document_text)} chars")
            
            # Check if we have anything to process
            if not script_text and not document_text:
                logger.warning(f"⚠️ No script or documents found for agent {agent_id}")
                return {
                    "success": False,
                    "error": "No script or documents to generate context from"
                }
            
            # Generate context using OpenAI
            summary_result = await openai_service.generate_agent_summary(
                document_text=document_text if document_text else None,
                script_text=script_text if script_text else None,
                existing_context=existing_context if not force_regenerate else None
            )
            
            if not summary_result.get("success"):
                logger.error(f"❌ Failed to generate summary: {summary_result.get('error')}")
                return {
                    "success": False,
                    "error": summary_result.get("error", "Summary generation failed")
                }
            
            # Get the generated context
            context = summary_result.get("context", {})
            
            # Add metadata
            context["generated_at"] = datetime.utcnow().isoformat()
            context["script_included"] = bool(script_text)
            context["documents_included"] = bool(document_text)
            
            # Get source document names
            source_docs = await db.agent_documents.find(
                {"agent_id": agent_id, "processing_status": "completed"},
                {"filename": 1}
            ).to_list(length=None)
            context["source_documents"] = [doc.get("filename", "") for doc in source_docs]
            
            # Store context in agent record
            await db.voice_agents.update_one(
                {"_id": ObjectId(agent_id)},
                {
                    "$set": {
                        "agent_context": context,
                        "has_context": True,
                        "context_generated_at": datetime.utcnow(),
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"✅ Agent context generated and stored successfully")
            logger.info(f"   Identity: {context.get('identity', {})}")
            logger.info(f"   Services: {len(context.get('company_info', {}).get('services', []))} items")
            logger.info(f"   FAQs: {len(context.get('faqs', []))} items")
            logger.info(f"   Source documents: {context.get('source_documents', [])}")
            
            return {
                "success": True,
                "context": context,
                "cached": False
            }
            
        except Exception as e:
            logger.error(f"❌ Error generating agent context: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e)
            }
    
    
    async def _get_all_document_text(
        self,
        agent_id: str,
        db: AsyncIOMotorDatabase
    ) -> str:
        """
        Get all extracted text from agent's documents
        
        Args:
            agent_id: The agent's ID
            db: Database connection
            
        Returns:
            Combined text from all documents
        """
        try:
            # Get all completed documents for this agent
            documents = await db.agent_documents.find({
                "agent_id": agent_id,
                "processing_status": "completed"
            }).to_list(length=None)
            
            if not documents:
                logger.info(f"ℹ️ No documents found for agent {agent_id}")
                return ""
            
            # Combine all document texts
            all_texts = []
            for doc in documents:
                extracted_text = doc.get("extracted_text", "")
                if extracted_text:
                    filename = doc.get("filename", "Unknown")
                    all_texts.append(f"--- Document: {filename} ---\n{extracted_text}")
            
            combined_text = "\n\n".join(all_texts)
            logger.info(f"📄 Combined {len(documents)} documents ({len(combined_text)} chars)")
            
            return combined_text
            
        except Exception as e:
            logger.error(f"❌ Error getting document text: {e}")
            return ""
    
    
    async def update_agent_context_on_script_change(
        self,
        agent_id: str,
        user_id: str,
        new_script: str,
        db: AsyncIOMotorDatabase
    ) -> Dict[str, Any]:
        """
        ✅ NEW: Update agent context when script changes
        
        Args:
            agent_id: The agent's ID
            user_id: The user's ID
            new_script: The new AI script
            db: Database connection
            
        Returns:
            Dict with success status
        """
        try:
            logger.info(f"🔄 Updating agent context for script change (Agent: {agent_id})")
            
            # Regenerate context with new script
            result = await self.generate_agent_context(
                agent_id=agent_id,
                user_id=user_id,
                db=db,
                script_text=new_script,
                force_regenerate=True
            )
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error updating context on script change: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    
    async def update_agent_context_on_document_change(
        self,
        agent_id: str,
        user_id: str,
        db: AsyncIOMotorDatabase
    ) -> Dict[str, Any]:
        """Update agent context when documents change (upload/delete)"""
        try:
            logger.info(f"🔄 Updating agent context for document change (Agent: {agent_id})")
            
            # Get agent first
            agent = await db.voice_agents.find_one({"_id": ObjectId(agent_id)})
            if not agent:
                return {"success": False, "error": "Agent not found"}
            
            # Regenerate context with current documents
            result = await self.generate_agent_context(
                agent_id=agent_id,
                user_id=user_id,
                db=db,
                force_regenerate=True
            )
            
            # ✅ NEW: Update greeting if context updated successfully
            if result.get("success"):
                context = result.get("context", {})
                identity = context.get("identity", {})
                company_name = identity.get("company")
                agent_name = agent.get("name", "AI Assistant")
                
                # Generate new greeting
                if company_name and company_name != "our company":
                    updated_greeting = f"Hi! This is {agent_name} from {company_name}. Thanks for taking my call."
                else:
                    updated_greeting = f"Hi! This is {agent_name}. Thanks for taking my call."
                
                # Update greeting in database
                await db.voice_agents.update_one(
                    {"_id": ObjectId(agent_id)},
                    {"$set": {
                        "greeting_message": updated_greeting,
                        "updated_at": datetime.utcnow()
                    }}
                )
                
                logger.info(f"✅ Updated greeting after document change: {updated_greeting}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error updating context on document change: {e}")
            return {"success": False, "error": str(e)}
    
    
    async def get_agent_context(
        self,
        agent_id: str,
        db: AsyncIOMotorDatabase
    ) -> Optional[Dict[str, Any]]:
        """
        ✅ NEW: Get pre-generated agent context (fast lookup)
        
        Args:
            agent_id: The agent's ID
            db: Database connection
            
        Returns:
            Agent context dict or None
        """
        try:
            agent = await db.voice_agents.find_one(
                {"_id": ObjectId(agent_id)},
                {"agent_context": 1, "has_context": 1}
            )
            
            if agent and agent.get("has_context") and agent.get("agent_context"):
                return agent.get("agent_context")
            
            return None
            
        except Exception as e:
            logger.error(f"❌ Error getting agent context: {e}")
            return None
    
    
    # ============================================
    # DOCUMENT UPLOAD & PROCESSING
    # ============================================
    
    async def upload_and_process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        agent_id: str,
        user_id: str,
        db: AsyncIOMotorDatabase
    ) -> Dict[str, Any]:
        """
        Upload document and process it for RAG
        
        ✅ ENHANCED: Now also triggers context regeneration after upload
        
        Steps:
        1. Validate file
        2. Save to disk
        3. Extract text
        4. Create chunks
        5. Generate embeddings
        6. Store in database
        7. ✅ NEW: Regenerate agent context
        """
        try:
            logger.info(f"📄 Starting document upload for agent {agent_id}")
            
            # Step 1: Validate file
            validation = self._validate_file(filename, content_type, len(file_content))
            if not validation['valid']:
                return {
                    "success": False,
                    "error": validation['error']
                }
            
            # Step 2: Save file to disk
            file_path = await self._save_file(
                file_content,
                filename,
                agent_id,
                user_id
            )
            
            # Step 3: Create document record
            doc_id = await self._create_document_record(
                agent_id=agent_id,
                user_id=user_id,
                filename=filename,
                file_path=str(file_path),
                file_type=content_type,
                file_size=len(file_content),
                db=db
            )
            
            # Step 4: Extract text from document
            logger.info(f"📖 Extracting text from {filename}")
            extracted_text = await self._extract_text(
                file_content,
                content_type
            )
            
            if not extracted_text:
                await self._update_document_status(
                    doc_id,
                    "failed",
                    "Failed to extract text from document",
                    db
                )
                return {
                    "success": False,
                    "error": "Failed to extract text from document"
                }
            
            # Step 5: Create text chunks
            logger.info(f"✂️ Creating text chunks")
            chunks = self._create_chunks(extracted_text)
            
            # Step 6: Generate embeddings for chunks
            logger.info(f"🧠 Generating embeddings for {len(chunks)} chunks")
            chunks_with_embeddings = await self._generate_embeddings(chunks)
            
            # Step 7: Store chunks with embeddings in database
            await self._store_chunks(
                doc_id=doc_id,
                agent_id=agent_id,
                chunks=chunks_with_embeddings,
                extracted_text=extracted_text,
                db=db
            )
            
            # Step 8: Update agent's has_training_docs flag
            await self._update_agent_training_status(agent_id, doc_id, db)
            
            logger.info(f"✅ Document processed successfully: {filename}")
            
            # ✅ NEW: Step 9 - Regenerate agent context with new document
            logger.info(f"🔄 Regenerating agent context after document upload...")
            context_result = await self.update_agent_context_on_document_change(
                agent_id=agent_id,
                user_id=user_id,
                db=db
            )
            
            if context_result.get("success"):
                logger.info(f"✅ Agent context updated with new document")
            else:
                logger.warning(f"⚠️ Failed to update context: {context_result.get('error')}")
            
            return {
                "success": True,
                "document_id": doc_id,
                "filename": filename,
                "file_size": len(file_content),
                "total_chunks": len(chunks_with_embeddings),
                "processing_status": "completed",
                "context_updated": context_result.get("success", False)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing document: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }
    
    
    # ============================================
    # TEXT EXTRACTION
    # ============================================
    
    async def _extract_text(
        self,
        file_content: bytes,
        content_type: str
    ) -> Optional[str]:
        """Extract text from different file types"""
        try:
            if content_type == "application/pdf":
                return await self._extract_text_from_pdf(file_content)
            
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                return await self._extract_text_from_docx(file_content)
            
            elif content_type in ["text/plain", "text/markdown"]:
                return file_content.decode('utf-8', errors='ignore')
            
            else:
                logger.warning(f"Unsupported content type: {content_type}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return None
    
    
    async def _extract_text_from_pdf(self, file_content: bytes) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    
    async def _extract_text_from_docx(self, file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    
    # ============================================
    # TEXT CHUNKING
    # ============================================
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
        
        # Clean text
        text = text.strip()
        
        # Split into sentences first
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(" ".join(current_chunk))
                
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - 2)
                current_chunk = current_chunk[overlap_start:]
                current_length = sum(len(s) for s in current_chunk)
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Don't forget the last chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        logger.info(f"Created {len(chunks)} chunks from {len(text)} chars")
        return chunks
    
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        import re
        
        # Simple sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    
    # ============================================
    # EMBEDDING GENERATION
    # ============================================
    
    async def _generate_embeddings(
        self,
        chunks: List[str]
    ) -> List[Dict[str, Any]]:
        """Generate OpenAI embeddings for text chunks"""
        try:
            chunks_with_embeddings = []
            
            # Process chunks in batches to avoid rate limits
            batch_size = 10
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                
                # Call OpenAI API
                response = await self.openai_client.embeddings.create(
                    model=self.embedding_model,
                    input=batch
                )
                
                # Extract embeddings
                for idx, chunk_text in enumerate(batch):
                    embedding = response.data[idx].embedding
                    
                    chunks_with_embeddings.append({
                        "chunk_id": f"chunk_{i + idx}",
                        "text": chunk_text,
                        "embedding": embedding,
                        "chunk_index": i + idx
                    })
            
            logger.info(f"Generated {len(chunks_with_embeddings)} embeddings")
            return chunks_with_embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise
    
    
    # ============================================
    # SIMILARITY SEARCH (KEPT FOR BACKWARD COMPATIBILITY)
    # ============================================
    
    async def query_documents(
        self,
        agent_id: str,
        query: str,
        db: AsyncIOMotorDatabase,
        top_k: int = 3
    ) -> Dict[str, Any]:
        """
        Search agent's documents for relevant information
        
        ✅ NOTE: This is kept for backward compatibility.
        The new approach uses pre-generated context instead of real-time search.
        
        Returns:
        - relevant chunks
        - confidence score
        - combined answer
        """
        try:
            logger.info(f"\n{'='*60}")
            logger.info(f"🔍 RAG DOCUMENT SEARCH")
            logger.info(f"{'='*60}")
            logger.info(f"   Agent ID: {agent_id}")
            logger.info(f"   Query: '{query}'")
            logger.info(f"{'='*60}\n")
            
            # Step 1: Get agent's documents
            documents = await db.agent_documents.find({
                "agent_id": agent_id,
                "processing_status": "completed"
            }).to_list(length=None)
            
            logger.info(f"📄 Found {len(documents)} documents for agent {agent_id}")
            
            if not documents:
                # Check if documents exist with different status
                all_docs = await db.agent_documents.find({
                    "agent_id": agent_id
                }).to_list(length=None)
                
                if all_docs:
                    logger.warning(f"⚠️ Agent has {len(all_docs)} documents but none are 'completed'")
                    for doc in all_docs:
                        logger.warning(f"   - {doc.get('filename')}: status={doc.get('processing_status')}")
                else:
                    logger.info(f"ℹ️ No documents found for agent {agent_id}")
                
                return {
                    "found": False,
                    "confidence": 0.0,
                    "answer": None,
                    "sources": []
                }
            
            # Log document details
            for doc in documents:
                chunk_count = len(doc.get("chunks", []))
                logger.info(f"   📄 {doc.get('filename')}: {chunk_count} chunks")
            
            # Step 2: Generate query embedding
            logger.info(f"🧠 Generating query embedding...")
            query_embedding = await self._generate_query_embedding(query)
            logger.info(f"✅ Query embedding generated ({len(query_embedding)} dimensions)")
            
            # Step 3: Search all chunks for similarity
            all_matches = []
            search_threshold = self.similarity_threshold - 0.1  # Lower threshold for voice
            
            for doc in documents:
                chunks = doc.get("chunks", [])
                for chunk in chunks:
                    chunk_embedding = chunk.get("embedding", [])
                    if chunk_embedding:
                        similarity = self._cosine_similarity(query_embedding, chunk_embedding)
                        
                        if similarity >= search_threshold:
                            all_matches.append({
                                "text": chunk.get("text", ""),
                                "similarity": similarity,
                                "document_id": str(doc["_id"]),
                                "filename": doc.get("filename", "Unknown"),
                                "chunk_id": chunk.get("chunk_id", "")
                            })
            
            logger.info(f"📊 Found {len(all_matches)} matches above threshold {search_threshold}")
            
            # Step 4: Sort by similarity and get top K
            all_matches.sort(key=lambda x: x["similarity"], reverse=True)
            top_matches = all_matches[:top_k]
            
            if not top_matches:
                logger.info(f"❌ No relevant matches found above threshold {search_threshold}")
                return {
                    "found": False,
                    "confidence": 0.0,
                    "answer": None,
                    "sources": []
                }
            
            # Step 5: Combine top matches into answer
            avg_confidence = sum(m["similarity"] for m in top_matches) / len(top_matches)
            combined_text = "\n\n".join([m["text"] for m in top_matches])
            
            logger.info(f"\n✅ RAG RESULTS:")
            logger.info(f"   Top matches: {len(top_matches)}")
            logger.info(f"   Avg confidence: {avg_confidence:.2f}")
            for i, match in enumerate(top_matches):
                logger.info(f"   Match {i+1}: {match['similarity']:.3f} from {match['filename']}")
                logger.info(f"            '{match['text'][:80]}...'")
            
            # Step 6: Update usage stats
            for match in top_matches:
                await db.agent_documents.update_one(
                    {"_id": ObjectId(match["document_id"])},
                    {
                        "$inc": {"usage_count": 1},
                        "$set": {"last_used": datetime.utcnow()}
                    }
                )
            
            return {
                "found": True,
                "confidence": avg_confidence,
                "answer": combined_text,
                "sources": top_matches,
                "total_matches": len(all_matches)
            }
            
        except Exception as e:
            logger.error(f"❌ Error querying documents: {e}", exc_info=True)
            return {
                "found": False,
                "confidence": 0.0,
                "answer": None,
                "error": str(e)
            }
    
    
    async def _generate_query_embedding(self, query: str) -> List[float]:
        """Generate embedding for search query"""
        try:
            response = await self.openai_client.embeddings.create(
                model=self.embedding_model,
                input=[query]
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Error generating query embedding: {e}")
            raise
    
    
    def _cosine_similarity(
        self,
        vec1: List[float],
        vec2: List[float]
    ) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            vec1_np = np.array(vec1)
            vec2_np = np.array(vec2)
            
            dot_product = np.dot(vec1_np, vec2_np)
            norm1 = np.linalg.norm(vec1_np)
            norm2 = np.linalg.norm(vec2_np)
            
            if norm1 == 0 or norm2 == 0:
                return 0.0
            
            return float(dot_product / (norm1 * norm2))
        except Exception as e:
            logger.error(f"Error calculating cosine similarity: {e}")
            return 0.0
    
    
    # ============================================
    # FILE MANAGEMENT
    # ============================================
    
    def _validate_file(
        self,
        filename: str,
        content_type: str,
        file_size: int
    ) -> Dict[str, Any]:
        """Validate uploaded file"""
        
        # Check file size (max 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if file_size > max_size:
            return {
                "valid": False,
                "error": f"File size exceeds maximum allowed size of 10MB (got {file_size / 1024 / 1024:.2f}MB)"
            }
        
        # Check file type
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain",
            "text/markdown"
        ]
        
        if content_type not in allowed_types:
            return {
                "valid": False,
                "error": f"Unsupported file type: {content_type}. Allowed: PDF, DOCX, TXT, MD"
            }
        
        return {"valid": True}
    
    
    async def _save_file(
        self,
        file_content: bytes,
        filename: str,
        agent_id: str,
        user_id: str
    ) -> Path:
        """Save uploaded file to disk"""
        try:
            # Create directory structure: uploads/agent_documents/{user_id}/{agent_id}/
            agent_dir = self.upload_base_path / user_id / agent_id
            agent_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate unique filename with hash
            file_hash = hashlib.md5(file_content).hexdigest()[:8]
            safe_filename = f"{file_hash}_{filename}"
            
            file_path = agent_dir / safe_filename
            
            # Write file asynchronously
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"💾 Saved file to {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise
    
    
    async def delete_document(
        self,
        document_id: str,
        agent_id: str,
        user_id: str,
        db: AsyncIOMotorDatabase
    ) -> Dict[str, Any]:
        """
        Delete document and its file
        
        ✅ ENHANCED: Now also triggers context regeneration after delete
        """
        try:
            # Get document record
            doc = await db.agent_documents.find_one({
                "_id": ObjectId(document_id),
                "agent_id": agent_id,
                "user_id": user_id
            })
            
            if not doc:
                return {
                    "success": False,
                    "error": "Document not found"
                }
            
            # Delete file from disk
            file_path = Path(doc["file_path"])
            if file_path.exists():
                file_path.unlink()
                logger.info(f"🗑️ Deleted file: {file_path}")
            
            # Delete document record
            await db.agent_documents.delete_one({"_id": ObjectId(document_id)})
            
            # Update agent's training doc list
            await db.voice_agents.update_one(
                {"_id": ObjectId(agent_id)},
                {
                    "$pull": {"training_doc_ids": document_id},
                    "$set": {"updated_at": datetime.utcnow()}
                }
            )
            
            # Check if agent still has documents
            remaining_docs = await db.agent_documents.count_documents({
                "agent_id": agent_id,
                "processing_status": "completed"
            })
            
            if remaining_docs == 0:
                await db.voice_agents.update_one(
                    {"_id": ObjectId(agent_id)},
                    {"$set": {"has_training_docs": False}}
                )
            
            logger.info(f"✅ Document deleted: {document_id}")
            
            # ✅ NEW: Regenerate agent context after document deletion
            logger.info(f"🔄 Regenerating agent context after document deletion...")
            context_result = await self.update_agent_context_on_document_change(
                agent_id=agent_id,
                user_id=user_id,
                db=db
            )
            
            if context_result.get("success"):
                logger.info(f"✅ Agent context updated after document deletion")
            else:
                logger.warning(f"⚠️ Failed to update context: {context_result.get('error')}")
            
            return {
                "success": True,
                "message": "Document deleted successfully",
                "context_updated": context_result.get("success", False)
            }
            
        except Exception as e:
            logger.error(f"❌ Error deleting document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    
    async def _create_document_record(
        self,
        agent_id: str,
        user_id: str,
        filename: str,
        file_path: str,
        file_type: str,
        file_size: int,
        db: AsyncIOMotorDatabase
    ) -> str:
        """Create initial document record in database"""
        doc_data = {
            "agent_id": agent_id,
            "user_id": user_id,
            "filename": filename,
            "file_path": file_path,
            "file_type": file_type,
            "file_size": file_size,
            "processing_status": "processing",
            "processing_started_at": datetime.utcnow(),
            "chunks": [],
            "total_chunks": 0,
            "usage_count": 0,
            "upload_date": datetime.utcnow(),
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        result = await db.agent_documents.insert_one(doc_data)
        return str(result.inserted_id)
    
    
    async def _store_chunks(
        self,
        doc_id: str,
        agent_id: str,
        chunks: List[Dict[str, Any]],
        extracted_text: str,
        db: AsyncIOMotorDatabase
    ):
        """Store processed chunks in database"""
        await db.agent_documents.update_one(
            {"_id": ObjectId(doc_id)},
            {
                "$set": {
                    "chunks": chunks,
                    "total_chunks": len(chunks),
                    "extracted_text": extracted_text[:10000],  # Store first 10K chars
                    "processing_status": "completed",
                    "processing_completed_at": datetime.utcnow(),
                    "updated_at": datetime.utcnow(),
                    "metadata": {
                        "word_count": len(extracted_text.split()),
                        "char_count": len(extracted_text)
                    }
                }
            }
        )
    
    
    async def _update_document_status(
        self,
        doc_id: str,
        status: str,
        error_message: Optional[str],
        db: AsyncIOMotorDatabase
    ):
        """Update document processing status"""
        update_data = {
            "processing_status": status,
            "updated_at": datetime.utcnow()
        }
        
        if error_message:
            update_data["processing_error"] = error_message
        
        await db.agent_documents.update_one(
            {"_id": ObjectId(doc_id)},
            {"$set": update_data}
        )
    
    
    async def _update_agent_training_status(
        self,
        agent_id: str,
        doc_id: str,
        db: AsyncIOMotorDatabase
    ):
        """Update agent's training document flags"""
        await db.voice_agents.update_one(
            {"_id": ObjectId(agent_id)},
            {
                "$set": {
                    "has_training_docs": True,
                    "updated_at": datetime.utcnow()
                },
                "$addToSet": {
                    "training_doc_ids": doc_id
                }
            }
        )
    
    
    # ============================================
    # UTILITY METHODS
    # ============================================
    
    async def get_agent_documents(
        self,
        agent_id: str,
        user_id: str,
        db: AsyncIOMotorDatabase
    ) -> List[Dict[str, Any]]:
        """Get all documents for an agent"""
        try:
            documents = await db.agent_documents.find({
                "agent_id": agent_id,
                "user_id": user_id
            }).to_list(length=None)
            
            # Convert ObjectId to string
            for doc in documents:
                doc["_id"] = str(doc["_id"])
                doc["upload_date"] = doc["upload_date"].isoformat()
                
                # Remove embeddings from response (too large)
                if "chunks" in doc:
                    for chunk in doc["chunks"]:
                        chunk.pop("embedding", None)
            
            return documents
            
        except Exception as e:
            logger.error(f"Error getting agent documents: {e}")
            return []


# Create singleton instance
rag_service = RAGService()